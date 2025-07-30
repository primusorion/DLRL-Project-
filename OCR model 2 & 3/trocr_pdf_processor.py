import os
import datetime
from pdf2image import convert_from_path
from PIL import Image, ImageOps
from transformers import TrOCRProcessor, VisionEncoderDecoderModel
import torch
import pdfplumber
from docx import Document
from fpdf import FPDF

# Set up
PDF_PATH = "China_Janes_Fighting_Ships_2023-2024_20250616_160038.pdf"
POPDIR = r'D:/propeller/poppler-24.08.0/Library/bin'
OUTPUT_FOLDER = "outputs"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Load TrOCR model
processor = TrOCRProcessor.from_pretrained("microsoft/trocr-base-printed")
model = VisionEncoderDecoderModel.from_pretrained("microsoft/trocr-base-printed")
device = "cuda" if torch.cuda.is_available() else "cpu"
model.to(device)

def extract_text_from_image(img):
    img = img.convert("RGB")
    img = img.resize((img.width * 2, img.height * 2))  # upscale for better OCR
    pixel_values = processor(images=img, return_tensors="pt").pixel_values.to(device)
    generated_ids = model.generate(pixel_values)
    return processor.batch_decode(generated_ids, skip_special_tokens=True)[0]

def extract_text_and_images_from_pdf(pdf_path):
    text_output = ""
    embedded_images = []

    # OCR page text
    pages = convert_from_path(pdf_path, dpi=300, poppler_path=POPDIR)
    for idx, page_img in enumerate(pages):
        print(f"🔍 OCR Page {idx+1}")
        page_text = extract_text_from_image(page_img)
        text_output += f"--- Page {idx+1} ---\n{page_text}\n\n"
        img_path = os.path.join(OUTPUT_FOLDER, f"page_{idx+1}.png")
        page_img.save(img_path)

    # Extract embedded images separately
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            im = page.to_image(resolution=300).original.convert("RGB")
            for j, obj in enumerate(page.images):
                x0, top, x1, bottom = obj['x0'], obj['top'], obj['x1'], obj['bottom']
                h = page.height
                img_crop = im.crop((x0, h - bottom, x1, h - top))
                out_path = os.path.join(OUTPUT_FOLDER, f"embedded_page{i+1}_img{j+1}.png")
                img_crop.save(out_path)
                embedded_images.append(out_path)

    return text_output, embedded_images

def save_to_docx(text, images, filename):
    doc = Document()
    for para in text.strip().split('\n'):
        doc.add_paragraph(para)
    for img in images:
        doc.add_paragraph()
        doc.add_picture(img, width=None)
    doc.save(filename)

def save_to_pdf(text, images, filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.set_auto_page_break(auto=True, margin=15)
    for line in text.strip().split('\n'):
        pdf.multi_cell(0, 10, line)
    for img in images:
        pdf.add_page()
        pdf.image(img, x=10, y=10, w=180)
    pdf.output(filename)

# Run extraction
print("📄 Processing PDF...")
text, images = extract_text_and_images_from_pdf(PDF_PATH)

timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
output_docx = os.path.join(OUTPUT_FOLDER, f"output_{timestamp}.docx")
output_pdf = os.path.join(OUTPUT_FOLDER, f"output_{timestamp}.pdf")

print("💾 Saving DOCX...")
save_to_docx(text, images, output_docx)

print("💾 Saving PDF...")
save_to_pdf(text, images, output_pdf)

print("✅ Done! Files saved in:", OUTPUT_FOLDER)
