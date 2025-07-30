import os
import re
import datetime
import unicodedata
import pytesseract
import cv2
import numpy as np
from PIL import Image, ImageOps, UnidentifiedImageError
from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename
from fpdf import FPDF
from docx import Document
from docx import Document as DocReader
from pdf2image import convert_from_path
from transformers import TrOCRProcessor, VisionEncoderDecoderModel
import torch

# ========== Flask Setup ==========
app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
FONT_PATH = 'static/fonts/DejaVuSans.ttf'
POPDIR = r'D:/propeller/poppler-24.08.0/Library/bin'
pytesseract.pytesseract.tesseract_cmd = r'D:/tesseract/tesseract.exe'

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# ========== TrOCR Setup ==========
processor = TrOCRProcessor.from_pretrained("microsoft/trocr-base-printed")
model = VisionEncoderDecoderModel.from_pretrained("microsoft/trocr-base-printed")
device = "cuda" if torch.cuda.is_available() else "cpu"
model.to(device)

# ========== Helpers ==========
def clean_text(text):
    text = text.strip()
    text = unicodedata.normalize("NFKD", text)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text

def save_to_txt(text, filename):
    with open(filename, "w", encoding="utf-8") as f:
        f.write(text)

def save_to_pdf_with_image(text, image_paths, filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font('DejaVu', '', FONT_PATH, uni=True)
    pdf.set_font("DejaVu", size=12)
    pdf.set_auto_page_break(auto=True, margin=15)
    for line in text.split('\n'):
        pdf.multi_cell(0, 10, line)
    for img_path in image_paths:
        pdf.add_page()
        pdf.image(img_path, x=10, y=10, w=180)
    pdf.output(filename)

def save_to_docx_with_images(text, image_paths, filename):
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    for img_path in image_paths:
        doc.add_paragraph()
        doc.add_picture(img_path, width=None)
    doc.save(filename)

# ========== OCR Methods ==========
def extract_text_from_image_trocr(image_path):
    img = Image.open(image_path).convert("RGB")
    img = img.resize((img.width * 2, img.height * 2))
    pixel_values = processor(images=img, return_tensors="pt").pixel_values.to(device)
    generated_ids = model.generate(pixel_values)
    trocr_text = processor.batch_decode(generated_ids, skip_special_tokens=True)[0]
    return trocr_text + "\n\n[via TrOCR]"

def extract_text_from_image_tesseract(image_path, lang='eng'):
    img = Image.open(image_path).convert("L")
    img = ImageOps.invert(img)
    img = img.resize((img.width * 2, img.height * 2))
    text = pytesseract.image_to_string(img, lang=lang)
    return text + "\n\n[via Tesseract]"

def extract_text_and_images_from_pdf(pdf_path, lang='eng'):
    text = ""
    image_files = []
    pages = convert_from_path(pdf_path, dpi=300, poppler_path=POPDIR)

    for i, page in enumerate(pages):
        page_img_path = os.path.join(OUTPUT_FOLDER, f"page_{i+1}.png")
        page.save(page_img_path)

        gray = page.convert("L")
        inv = ImageOps.invert(gray)
        resized = inv.resize((inv.width * 2, inv.height * 2))
        page_text = pytesseract.image_to_string(resized, lang=lang)
        text += f"\n--- Page {i+1} [via Tesseract] ---\n{page_text}"

        # Submarine image detection
        img_cv = cv2.cvtColor(np.array(page.convert("RGB")), cv2.COLOR_RGB2BGR)
        gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
        blur = cv2.GaussianBlur(gray, (5, 5), 0)
        _, thresh = cv2.threshold(blur, 180, 255, cv2.THRESH_BINARY_INV)
        contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

        for j, cnt in enumerate(contours):
            x, y, w, h = cv2.boundingRect(cnt)
            if w > 100 and h > 100 and h / w < 3:
                crop = page.crop((x, y, x + w, y + h))
                img_path = os.path.join(OUTPUT_FOLDER, f"page{i+1}_img{j+1}.png")
                crop.save(img_path)
                image_files.append(img_path)

    return text.strip(), image_files

def extract_text_from_docx(docx_path):
    doc = DocReader(docx_path)
    text = "\n".join([p.text for p in doc.paragraphs])
    return text.strip() + "\n\n[via DOCX extract]", []

# ========== Flask Route ==========
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        file = request.files.get('file')
        if not file:
            return "No file uploaded. Please select a file and try again.", 400

        output_format = request.form.get('format')
        lang = request.form.get('lang', 'eng')

        filename = secure_filename(file.filename)
        input_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(input_path)

        base_name = os.path.splitext(filename)[0]
        ext = filename.lower().split('.')[-1]
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_name = f"{base_name}_{timestamp}.{output_format}"
        output_path = os.path.join(OUTPUT_FOLDER, output_name)

        try:
            if ext == 'pdf':
                text, image_paths = extract_text_and_images_from_pdf(input_path, lang)
            elif ext == 'docx':
                text, image_paths = extract_text_from_docx(input_path)
            elif ext in ['jpg', 'jpeg', 'png']:
                text = extract_text_from_image_trocr(input_path)
                image_paths = [input_path]
            else:
                return "Unsupported file format", 400

            text = clean_text(text)

            if output_format == 'txt':
                save_to_txt(text, output_path)
            elif output_format == 'pdf':
                save_to_pdf_with_image(text, image_paths, output_path)
            elif output_format == 'docx':
                save_to_docx_with_images(text, image_paths, output_path)
            else:
                return "Unsupported output format", 400

            return send_file(output_path, as_attachment=True, download_name=os.path.basename(output_path))


        except UnidentifiedImageError:
            return "The uploaded file is not a supported image or valid PDF.", 400
        except Exception as e:
            return f"Error: {str(e)}", 500

    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True)
