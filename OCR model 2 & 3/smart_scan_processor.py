import os
import datetime
import cv2
import numpy as np
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
from transformers import TrOCRProcessor, VisionEncoderDecoderModel
import torch
from docx import Document
from fpdf import FPDF

# ========== CONFIG ==========
PDF_PATH = "China_Janes_Fighting_Ships_2023-2024.pdf"
POPDIR = r"D:/propeller/poppler-24.08.0/Library/bin"
OUTPUT_FOLDER = "final_output"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# ✅ Tesseract executable path
pytesseract.pytesseract.tesseract_cmd = r'D:\tesseract\tesseract.exe'

# ========== Load TrOCR ==========
processor = TrOCRProcessor.from_pretrained("microsoft/trocr-base-printed")
model = VisionEncoderDecoderModel.from_pretrained("microsoft/trocr-base-printed")
device = "cuda" if torch.cuda.is_available() else "cpu"
model.to(device)

# ========== OCR with fallback ==========
def ocr_with_trocr_and_fallback(pil_img):
    # Preprocess
    img_gray = pil_img.convert("L").resize((pil_img.width * 2, pil_img.height * 2))
    img_eq = Image.fromarray(cv2.equalizeHist(np.array(img_gray))).convert("RGB")

    # TrOCR first
    try:
        pixel_values = processor(images=img_eq, return_tensors="pt").pixel_values.to(device)
        generated_ids = model.generate(pixel_values)
        trocr_text = processor.batch_decode(generated_ids, skip_special_tokens=True)[0].strip()
        if trocr_text:
            return trocr_text + "  [via TrOCR]"
    except Exception as e:
        print("TrOCR failed:", e)

    # Tesseract fallback
    try:
        text_fallback = pytesseract.image_to_string(img_gray)
        return text_fallback.strip() + "  [via Tesseract]"
    except Exception as e:
        return "[OCR Failed]"

# ========== Image + Text extraction ==========
def extract_text_and_images_from_page(pil_img, page_index):
    text = ocr_with_trocr_and_fallback(pil_img)
    extracted_images = []

    # Submarine image extraction (v1 logic)
    img_cv = np.array(pil_img.convert("RGB"))
    gray = cv2.cvtColor(img_cv, cv2.COLOR_RGB2GRAY)
    blur = cv2.GaussianBlur(gray, (5, 5), 0)
    _, thresh = cv2.threshold(blur, 180, 255, cv2.THRESH_BINARY_INV)
    contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    for i, cnt in enumerate(contours):
        x, y, w, h = cv2.boundingRect(cnt)
        if w > 100 and h > 100 and h / w < 3:
            crop = pil_img.crop((x, y, x + w, y + h))
            out_path = os.path.join(OUTPUT_FOLDER, f"page{page_index+1}_img{i+1}.png")
            crop.save(out_path)
            extracted_images.append(out_path)

    return text.strip(), extracted_images

# ========== PDF Processor ==========
def process_pdf(pdf_path):
    pages = convert_from_path(pdf_path, dpi=300, poppler_path=POPDIR)
    all_text = ""
    all_images = []

    for i, page_img in enumerate(pages):
        print(f"📄 Processing Page {i+1}")
        text, images = extract_text_and_images_from_page(page_img, i)
        all_text += f"\n--- Page {i+1} ---\n{text}\n"
        all_images.extend(images)

    return all_text.strip(), all_images

# ========== Output Writers ==========
def save_to_txt(text, path):
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)

def save_to_docx(text, images, path):
    doc = Document()
    for line in text.strip().split('\n'):
        doc.add_paragraph(line)
    for img in images:
        doc.add_paragraph()
        doc.add_picture(img, width=None)
    doc.save(path)

def save_to_pdf(text, images, path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in text.strip().split('\n'):
        pdf.multi_cell(0, 10, line)
    for img in images:
        pdf.add_page()
        pdf.image(img, x=10, y=10, w=180)
    pdf.output(path)

# ========== Run ==========
if __name__ == "__main__":
    print("🚀 Starting Smart OCR Scan...")
    text, images = process_pdf(PDF_PATH)

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    txt_path = os.path.join(OUTPUT_FOLDER, f"output_{timestamp}.txt")
    docx_path = os.path.join(OUTPUT_FOLDER, f"output_{timestamp}.docx")
    pdf_path = os.path.join(OUTPUT_FOLDER, f"output_{timestamp}.pdf")

    print("💾 Saving TXT...")
    save_to_txt(text, txt_path)

    print("💾 Saving DOCX...")
    save_to_docx(text, images, docx_path)

    print("💾 Saving PDF...")
    save_to_pdf(text, images, pdf_path)

    print("\n✅ Done!")
    print("📂 Text:", txt_path)
    print("📂 DOCX:", docx_path)
    print("📂 PDF: ", pdf_path)
